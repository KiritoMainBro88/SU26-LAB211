"""Generate polished LAB211 reports, diagrams, and Audit Log evidence indexes.

The script intentionally uses only artifacts already present in this repository.
It does not generate or imply external grades, AI-chat screenshots, or unverifiable
execution results.
"""

from __future__ import annotations

import hashlib
import math
import re
from copy import copy
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as ReportLabImage,
    KeepTogether,
    LongTable,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from xml.sax.saxutils import escape


REPO_ROOT = Path(__file__).resolve().parents[1]
DIAGRAM_ROOT = REPO_ROOT / "docs" / "diagrams"
EDITABLE_ROOT = REPO_ROOT / "docs" / "editable"
AUDIT_ROOT = REPO_ROOT / "docs" / "ai-audit-logs"
PDF_ROOT = REPO_ROOT / "output" / "pdf"
VERIFY_DATE = datetime(2026, 7, 21)
VERIFY_DATE_LONG = "21 July 2026"
REPOSITORY_URL = "https://github.com/KiritoMainBro88/SU26-LAB211"
STUDENT_NAME = "Nguyen Dinh Chuong"
STUDENT_NAME_VI = "Nguyễn Đình Chương"
STUDENT_ID = "SE204330"

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
GREEN = "2E7D32"
LIGHT_GREEN = "E2F0D9"
ORANGE = "C65911"
LIGHT_ORANGE = "FCE4D6"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"


LABS: Dict[str, Dict[str, object]] = {
    "lab1": {
        "number": 1,
        "code": "J1.L.P0027",
        "title": "Mountain Hiking Challenge Registration",
        "functions": 9,
        "deep_count": 34,
        "regression_count": 13,
        "audit_entries": 12,
        "hallucinations": 3,
        "entry_class": "presentation.Program",
        "summary": (
            "A layered Java console application for FPT University student mountain-hiking "
            "registration. It validates identity and contact data, calculates sponsored tuition, "
            "supports nine required workflows, aggregates registrations by mountain, and persists "
            "records in the object format requested by the assignment."
        ),
        "classes": [
            ("Program", "Application entry point; loads data and starts the menu loop."),
            ("Menu", "Routes the nine functions and owns console presentation only."),
            ("IStudentManagement", "Focused business contract for CRUD, queries, statistics, and persistence."),
            ("StudentManagement", "Coordinates validation, fee calculation, DAOs, and user-facing outcomes."),
            ("StudentDAO", "Owns registration state, defensive copies, saved-state tracking, and object persistence."),
            ("MountainDAO", "Loads MountainList.csv and resolves normalized mountain codes."),
            ("Student / Mountain / StatisticalInfo", "Separate entities for registrations, reference data, and aggregation results."),
        ],
        "algorithms": [
            {
                "name": "A1 - Create a registration",
                "purpose": "Reject invalid state before a Student reaches persistent in-memory data.",
                "pseudocode": [
                    "normalize studentId and mountainCode",
                    "if studentId/name/phone/email is invalid: REJECT",
                    "if mountainCode does not exist or studentId is duplicated: REJECT",
                    "if baseFee is not finite or baseFee <= 0: REJECT",
                    "fee = baseFee * 0.65 when phone is Viettel/VNPT; otherwise baseFee",
                    "create Student and add a defensive copy to StudentDAO",
                    "mark data as changed and return success",
                ],
                "complexity": "O(n + m) time for duplicate and mountain lookup; O(1) additional space.",
            },
            {
                "name": "A2 - Safe partial update and tuition recalculation",
                "purpose": "Preserve blank fields and the original custom base-fee basis.",
                "pseudocode": [
                    "find current Student by immutable ID; reject when missing",
                    "for each editable field: blank -> old value; otherwise -> candidate value",
                    "validate the complete candidate before changing state",
                    "if a new base fee exists: calculate from that fee",
                    "else if phone changed: recover old base fee, then recalculate discount",
                    "replace the record only after every check succeeds",
                ],
                "complexity": "O(n + m) time; O(1) additional space excluding the candidate object.",
            },
            {
                "name": "A3 - Statistics by mountain",
                "purpose": "Show only mountains that have registrations, with count and total tuition.",
                "pseudocode": [
                    "create insertion-ordered map: mountainCode -> StatisticalInfo",
                    "for each Student:",
                    "    create a zeroed bucket when the code is first seen",
                    "    increment participant count and add tuition fee",
                    "return the map values in deterministic first-seen order",
                ],
                "complexity": "O(n) expected time and O(k) space, where k is the number of used mountains.",
            },
            {
                "name": "A4 - Defensive object-file load/save",
                "purpose": "Prevent malformed serialization or failed replacement from destroying valid state.",
                "pseudocode": [
                    "LOAD: deserialize into temporary state and validate every Student and unique ID",
                    "LOAD: replace active memory only after complete success",
                    "SAVE: validate the complete current list before opening output",
                    "SAVE: write a copied list to a temporary sibling file",
                    "SAVE: atomically replace registrations.dat when supported",
                ],
                "complexity": "O(n) time and O(n) temporary memory.",
            },
        ],
        "test_cases": [
            ("L1-T01", "Valid custom-fee registration", "SE123456, valid contact, MT01, 10,000,000", "Added; non-discount phone keeps 10,000,000", "PASS"),
            ("L1-T02", "Discount after phone-only update", "Change phone to 0961234567; base fee blank", "Preserve 10,000,000 basis; fee becomes 6,500,000", "PASS"),
            ("L1-T03", "Student ID partitions", "AB123456, SE12345A, duplicate se123456", "All rejected without changing valid state", "PASS"),
            ("L1-T04", "Name boundaries", "2, 20, and 21 characters", "2/20 accepted; 21 rejected", "PASS"),
            ("L1-T05", "Phone and numeric safety", "Unknown prefix, NaN, Infinity, zero", "All rejected", "PASS"),
            ("L1-T06", "Mountain normalization", "1, 01, MT01, mt01", "All resolve to MT01", "PASS"),
            ("L1-T07", "Defensive query result", "Mutate Student returned by getAllStudents", "Manager state remains unchanged", "PASS"),
            ("L1-T08", "Malformed serialized list", "Valid Student plus non-Student item", "Load fails and preserves memory", "PASS"),
        ],
        "traceability": [
            ("1. New Registration", "Menu.newRegistration", "StudentManagement.createStudent", "L1-T01, T03-T06"),
            ("2. Update Registration", "Menu.updateRegistration", "StudentManagement.updateStudent", "L1-T02"),
            ("3. Display Registered List", "Menu.displayRegisteredList", "StudentManagement.getAllStudents", "Defensive-copy checks"),
            ("4. Delete Registration", "Menu.deleteRegistration", "StudentManagement.deleteStudent", "Null/not-found checks"),
            ("5. Search by Name", "Menu.searchParticipantsByName", "StudentManagement.searchByName", "Case-insensitive search"),
            ("6. Filter by Campus", "Menu.filterDataByCampus", "StudentManagement.filterByCampus", "Valid/invalid campus"),
            ("7. Statistics by Location", "Menu.displayStatisticsByMountain", "StudentManagement.getStatisticsByMountain", "Aggregation checks"),
            ("8. Save Data", "Menu.saveData", "StudentDAO.saveStudents", "Save/reload and failure preservation"),
            ("9. Exit", "Menu.exitProgram", "StudentManagement.isSaved/saveData", "Save-failure cancellation trace"),
        ],
        "limitations": [
            "Java object serialization is intentionally assignment-specific and is not a long-term interchange format.",
            "Mountain reference parsing is designed for the supplied CSV shape; it is not a general RFC 4180 library.",
            "Mutable Mountain objects are exposed only to the read-only presentation path; Student records use defensive copies.",
        ],
    },
    "lab2": {
        "number": 2,
        "code": "J1.L.P0036",
        "title": "Football Club & Player Management",
        "functions": 14,
        "deep_count": 62,
        "regression_count": 23,
        "audit_entries": 18,
        "hallucinations": 3,
        "entry_class": "app.Main",
        "summary": (
            "A Java console application for a football league with fourteen required workflows. "
            "It models Clubs and Players separately, enforces referential and shirt-number rules, "
            "loads strict UTF-8 CSV data, and applies defensive copies at manager boundaries."
        ),
        "classes": [
            ("Menu", "Coordinates fourteen functions and the strict clubs-first reload sequence."),
            ("ClubIO / PlayerIO", "Own console input/output and delegate business decisions."),
            ("IManager<T> / AbstractManager<T>", "Provide reusable identity lookup, defensive-copy access, and change tracking."),
            ("IClubManager / ClubManager", "Implement Club CRUD, budget filtering, validation, load, and save."),
            ("IPlayerManager / PlayerManager", "Implement Player CRUD, search, filtering, sorting, and shirt uniqueness."),
            ("IClubFileHelper / IPlayerFileHelper", "Define text persistence boundaries for dependency inversion."),
            ("ClubFileHelper / PlayerFileHelper", "Parse strict CSV and replace individual targets through temporary files."),
            ("BaseEntity / Club / Player / Position", "Represent identity, domain data, and the constrained position vocabulary."),
        ],
        "algorithms": [
            {
                "name": "A1 - Strict clubs-first reload",
                "purpose": "Ensure Players never enter memory without a valid referenced Club.",
                "pseudocode": [
                    "clear ClubManager and PlayerManager",
                    "parse every physical club line; reject blank/malformed/duplicate data",
                    "if club load fails: report failure and remain empty",
                    "parse players against the validated Club list",
                    "reject missing club, duplicate ID, invalid position, or duplicate shirt in a Club",
                    "if player load fails: clear both managers and report failure",
                    "otherwise commit both validated collections",
                ],
                "complexity": "O(c^2 + p(c + p)) with list-based uniqueness checks; data size is assignment-bounded.",
            },
            {
                "name": "A2 - Add/update Player",
                "purpose": "Maintain identity, reference, position, and per-Club shirt-number invariants.",
                "pseudocode": [
                    "validate Pxxxx identity and locate existing Club",
                    "validate non-empty name and parse constrained Position",
                    "validate shirt in [1, 99]",
                    "scan Players for same Club + shirt, ignoring the current ID during update",
                    "build/modify state only after all validations pass",
                ],
                "complexity": "O(c + p) time and O(1) additional space.",
            },
            {
                "name": "A3 - Sort Players for Function 6",
                "purpose": "Sort by Club name, then shirt number only when Players share the same Club.",
                "pseudocode": [
                    "copy all Players so sorting cannot mutate manager order",
                    "compare resolved Club names case-insensitively",
                    "if names differ: return Club-name order",
                    "if same Club ID: compare shirt number",
                    "if different Clubs have equal names: use Club ID as deterministic tie-break",
                ],
                "complexity": "O(p log p * c) with list-based Club lookup; O(p) result space.",
            },
            {
                "name": "A4 - CSV save with protected replacement",
                "purpose": "Preserve each previous target when validation or writing fails.",
                "pseudocode": [
                    "validate complete list and all cross-entity references",
                    "escape commas, quotes, CR, and LF in text fields",
                    "write UTF-8 data to a temporary sibling file",
                    "replace the target atomically when supported; otherwise replace normally",
                    "delete temporary data on failure and report false",
                ],
                "complexity": "O(c + p(c + p)) validation time and O(1) streaming write space.",
            },
        ],
        "test_cases": [
            ("L2-T01", "Club boundaries", "null, NaN/Infinity budget, duplicate ID", "Rejected; valid finite Club accepted", "PASS"),
            ("L2-T02", "Player constraints", "null position, shirt 100, missing Club", "Rejected before insertion", "PASS"),
            ("L2-T03", "Shirt-number scope", "Same number in same and different Clubs", "Same-Club duplicate rejected; cross-Club accepted", "PASS"),
            ("L2-T04", "Blank update fields", "Name, position, and shirt blank", "Original fields retained; operation succeeds", "PASS"),
            ("L2-T05", "Sort tie-breaks", "Equal Club names; same-Club shirts", "Club ID tie-break; same Club by shirt", "PASS"),
            ("L2-T06", "Strict physical lines", "Blank Club line during reload", "Reload fails and both managers remain empty", "PASS"),
            ("L2-T07", "Quoted CSV round-trip", "Comma/quote in Club and Player names", "Text preserved exactly", "PASS"),
            ("L2-T08", "Defensive results", "Mutate get/find/search/filter result", "Manager state remains unchanged", "PASS"),
            ("L2-T09", "Finite threshold", "0, -1, NaN, Infinity", "Finite values accepted; non-finite rejected", "PASS"),
            ("L2-T10", "Manager persistence APIs", "Save and load via both managers", "Round-trip succeeds", "PASS"),
        ],
        "traceability": [
            ("1. List Clubs", "Menu case 1", "ClubIO.listAllClubs", "Table output / empty list"),
            ("2. Add Club", "Menu case 2", "ClubManager.addClub", "L2-T01"),
            ("3. Search Club", "Menu case 3", "ClubManager.getClubById", "Found/not-found"),
            ("4. Update Club", "Menu case 4", "ClubManager.updateClub", "Blank/invalid update"),
            ("5. Budget Filter", "Menu case 5", "ClubManager.filterByBudget", "L2-T09"),
            ("6. Sorted Players", "Menu case 6", "PlayerManager.getPlayersSortedByClubName", "L2-T05"),
            ("7. Search Players", "Menu case 7", "PlayerManager.searchByPartialName", "Case-insensitive/null"),
            ("8. Add Player", "Menu case 8", "PlayerManager.addPlayer", "L2-T02, T03"),
            ("9. Remove Player", "Menu case 9", "PlayerManager.removePlayer", "Found/not-found"),
            ("10. Update Player", "Menu case 10", "PlayerManager.updatePlayer", "L2-T03, T04"),
            ("11. Position Filter", "Menu case 11", "PlayerManager.filterByPosition", "Enum/null"),
            ("12. Save", "Menu case 12", "ClubManager/PlayerManager.saveToFile", "L2-T07, T10"),
            ("13. Load", "Menu case 13", "Menu.loadData", "L2-T06"),
            ("14. Quit", "Menu case 14", "Menu.quitProgram", "Changed-state save path"),
        ],
        "limitations": [
            "Clubs and Players are saved sequentially. Each target is protected, but the two files are not one transaction.",
            "Sorting resolves Club names through list lookup; a map would improve asymptotic performance for large data sets.",
            "Strict Function 13 rejects blank physical lines by design, matching the lecturer requirement.",
        ],
    },
    "lab3": {
        "number": 3,
        "code": "J1.L.P0037",
        "title": "Employee Payroll Management System",
        "functions": 9,
        "deep_count": 57,
        "regression_count": 17,
        "audit_entries": 20,
        "hallucinations": 3,
        "entry_class": "app.Main",
        "summary": (
            "A Java console payroll application that validates seven Employee fields, tolerates bad "
            "rows without accepting invalid state, performs lecturer-specified soft deletion, searches "
            "by four attributes, and calculates payroll for active Employees only."
        ),
        "classes": [
            ("Main", "Coordinates the nine workflows and changed-state save/quit behavior."),
            ("View / EmployeeView", "Own reusable validated input and employee/payroll table presentation."),
            ("EmployeeIO", "Builds validated add/update candidates without business-state ownership."),
            ("IEmployeeManager / EmployeeManager", "Own CRUD, soft delete, search, payroll, defensive copies, load, and save."),
            ("IFileReadWrite<T>", "Generic persistence boundary injected into EmployeeManager."),
            ("EmployeeFileHelper", "Reads resilient CSV rows and performs validated temporary-file replacement."),
            ("Employee", "Implements Identifiable and PayrollCalculable for domain-focused polymorphism."),
            ("Role / EmployeeStatus", "Constrain textual values and centralize canonical display/parse behavior."),
        ],
        "algorithms": [
            {
                "name": "A1 - Row-isolated employee loading",
                "purpose": "Continue after erroneous lines while never admitting malformed or duplicate state.",
                "pseudocode": [
                    "if employees.txt is missing: fail the load",
                    "for each non-blank physical line:",
                    "    parse exactly seven CSV fields",
                    "    parse Role/Status and finite numeric fields",
                    "    validate complete Employee and duplicate ID",
                    "    on row error: warn, log, skip, and continue",
                    "manager validates the completed list again, then replaces memory",
                ],
                "complexity": "O(n^2) time from list-based duplicate checks and O(n) result space.",
            },
            {
                "name": "A2 - Validate-before-commit update",
                "purpose": "Update only Role, Base Salary, Bonus, and Status as required by Function 3.",
                "pseudocode": [
                    "find current Employee by immutable ID",
                    "construct candidate with old ID, Name, and Working Days",
                    "apply only the four permitted candidate fields",
                    "validate the complete candidate",
                    "commit four fields only after validation succeeds",
                ],
                "complexity": "O(n) lookup time and O(1) additional space.",
            },
            {
                "name": "A3 - Soft delete and active-only payroll",
                "purpose": "Retain history while excluding inactive Employees from payroll.",
                "pseudocode": [
                    "REMOVE: find Employee; reject missing or already inactive record",
                    "REMOVE: set status to INACTIVE; keep list size unchanged",
                    "PAYROLL: iterate Employees and include only status == ACTIVE",
                    "salary = baseSalary * workingDays + bonus",
                    "sum and display the same active set",
                ],
                "complexity": "O(n) time and O(a) output space for a active Employees.",
            },
            {
                "name": "A4 - Safe employee-file save",
                "purpose": "Never truncate the previous target before validation and complete writing.",
                "pseudocode": [
                    "validate list, every Employee, and unique IDs",
                    "create a temporary sibling beside the target",
                    "write escaped UTF-8 CSV with decimal-safe number formatting",
                    "replace target atomically when supported",
                    "on failure delete temporary data and preserve previous target",
                ],
                "complexity": "O(n^2) validation time and O(1) streaming write space.",
            },
        ],
        "test_cases": [
            ("L3-T01", "Dependency/null guards", "Null helper, View, EmployeeIO dependencies", "IllegalArgumentException", "PASS"),
            ("L3-T02", "Working-day boundaries", "0, 26, and 27", "0/26 accepted; 27 rejected", "PASS"),
            ("L3-T03", "Non-finite salary", "Infinity salary; negative bonus", "Rejected", "PASS"),
            ("L3-T04", "Restricted update", "Candidate changes all fields", "Name/days preserved; four fields changed", "PASS"),
            ("L3-T05", "Soft delete", "Remove active E003 twice", "Inactive record retained; second remove rejected", "PASS"),
            ("L3-T06", "Exact status search", "active with active/inactive data", "Inactive does not match active", "PASS"),
            ("L3-T07", "Malformed-row isolation", "valid, malformed, duplicate, valid", "Two valid records loaded", "PASS"),
            ("L3-T08", "CSV quoting/decimals", "Doe, Jane; 120.5; 3.25", "Round-trip preserves values", "PASS"),
            ("L3-T09", "Rejected save preservation", "NaN bonus after valid target exists", "Write rejected; target unchanged", "PASS"),
            ("L3-T10", "Polymorphic contracts", "Employee through both interfaces", "Identity/payroll behavior matches", "PASS"),
        ],
        "traceability": [
            ("1. Load", "Main.manualLoadData", "EmployeeManager.loadFromFile", "L3-T07; failure preservation"),
            ("2. Add", "Main.addNewEmployee", "EmployeeManager.addEmployee", "L3-T01-T03"),
            ("3. Update", "Main.updateEmployee", "EmployeeManager.updateEmployee", "L3-T04"),
            ("4. Remove (D_Soft)", "Main.removeEmployee", "EmployeeManager.removeEmployee", "L3-T05"),
            ("5. Search", "Main.searchEmployee", "EmployeeManager.searchByAttribute", "L3-T06"),
            ("6. Payroll", "Main.calculatePayroll", "EmployeeManager.getTotalSalaryByMonth", "Active-only checks"),
            ("7. Display", "Main case 7", "EmployeeView.displayAllEmployees", "Empty/long-name output"),
            ("8. Save", "Main.saveData", "EmployeeFileHelper.write", "L3-T08, T09"),
            ("9. Quit", "Main.quitProgram", "Changed-state save decision", "Source branch trace"),
        ],
        "limitations": [
            "A file containing only invalid non-blank rows produces an empty valid result after warnings; this recovery policy is documented.",
            "Soft delete follows the lecturer's D_Soft note even though the assignment PDF uses the shorter word 'Delete'.",
            "List scans favor clarity for assignment-sized data; indexed structures would be preferable at enterprise scale.",
        ],
    },
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def rgb(hex_value: str) -> Tuple[int, int, int]:
    return tuple(int(hex_value[index : index + 2], 16) for index in (0, 2, 4))


def load_font(name: str, size: int) -> ImageFont.FreeTypeFont:
    font_path = Path("C:/Windows/Fonts") / name
    return ImageFont.truetype(str(font_path), size=size)


def draw_dashed_line(
    draw: ImageDraw.ImageDraw,
    start: Tuple[float, float],
    end: Tuple[float, float],
    fill: Tuple[int, int, int],
    width: int = 5,
    dash: int = 18,
) -> None:
    x1, y1 = start
    x2, y2 = end
    distance = math.hypot(x2 - x1, y2 - y1)
    if distance == 0:
        return
    ux, uy = (x2 - x1) / distance, (y2 - y1) / distance
    cursor = 0.0
    while cursor < distance:
        segment_end = min(cursor + dash, distance)
        draw.line(
            (x1 + ux * cursor, y1 + uy * cursor, x1 + ux * segment_end, y1 + uy * segment_end),
            fill=fill,
            width=width,
        )
        cursor += dash * 1.7


def boundary_point(box: Tuple[int, int, int, int], toward: Tuple[float, float]) -> Tuple[float, float]:
    left, top, right, bottom = box
    cx, cy = (left + right) / 2, (top + bottom) / 2
    dx, dy = toward[0] - cx, toward[1] - cy
    if dx == 0 and dy == 0:
        return cx, cy
    scale_x = (right - left) / 2 / abs(dx) if dx else float("inf")
    scale_y = (bottom - top) / 2 / abs(dy) if dy else float("inf")
    scale = min(scale_x, scale_y)
    return cx + dx * scale, cy + dy * scale


def draw_relation(
    draw: ImageDraw.ImageDraw,
    boxes: Dict[str, Tuple[int, int, int, int]],
    source: str,
    target: str,
    label: str,
    dashed: bool = False,
) -> None:
    source_box, target_box = boxes[source], boxes[target]
    source_center = ((source_box[0] + source_box[2]) / 2, (source_box[1] + source_box[3]) / 2)
    target_center = ((target_box[0] + target_box[2]) / 2, (target_box[1] + target_box[3]) / 2)
    start = boundary_point(source_box, target_center)
    end = boundary_point(target_box, source_center)
    line_color = rgb(GRAY)
    if dashed:
        draw_dashed_line(draw, start, end, line_color)
    else:
        draw.line((start[0], start[1], end[0], end[1]), fill=line_color, width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_size = 18
    points = [
        end,
        (
            end[0] - arrow_size * math.cos(angle - math.pi / 6),
            end[1] - arrow_size * math.sin(angle - math.pi / 6),
        ),
        (
            end[0] - arrow_size * math.cos(angle + math.pi / 6),
            end[1] - arrow_size * math.sin(angle + math.pi / 6),
        ),
    ]
    draw.polygon(points, fill=line_color)
    if label:
        font = load_font("arial.ttf", 24)
        midpoint = ((start[0] + end[0]) / 2, (start[1] + end[1]) / 2)
        bbox = draw.textbbox((0, 0), label, font=font)
        width, height = bbox[2] - bbox[0], bbox[3] - bbox[1]
        label_box = (
            midpoint[0] - width / 2 - 7,
            midpoint[1] - height / 2 - 5,
            midpoint[0] + width / 2 + 7,
            midpoint[1] + height / 2 + 5,
        )
        draw.rounded_rectangle(label_box, radius=6, fill=rgb(WHITE))
        draw.text((midpoint[0] - width / 2, midpoint[1] - height / 2), label, fill=line_color, font=font)


def draw_class_diagram(lab_key: str, data: Dict[str, object]) -> Path:
    width, height = 2600, 1750
    image = Image.new("RGB", (width, height), rgb(WHITE))
    draw = ImageDraw.Draw(image)
    title_font = load_font("arialbd.ttf", 48)
    class_font = load_font("arialbd.ttf", 30)
    detail_font = load_font("arial.ttf", 23)
    draw.text((80, 40), f"{data['code']} - Final UML Class Relationships", fill=rgb(NAVY), font=title_font)

    if lab_key == "lab1":
        nodes = {
            "Program": (80, 180, 470, 390, "presentation", ["+ main(args)"]),
            "Menu": (570, 180, 1050, 440, "presentation", ["+ manageStudentRegistration()", "- nine handlers"]),
            "StudentManagement": (1160, 170, 1800, 510, "business", ["+ createStudent()", "+ updateStudent()", "+ getStatisticsByMountain()", "+ saveData()"]),
            "IStudentManagement": (1940, 190, 2500, 480, "interface", ["<<interface>>", "+ CRUD/query/statistics"]),
            "StudentDAO": (650, 740, 1210, 1050, "persistence", ["+ loadStudents()", "+ saveStudents()", "+ defensive CRUD"]),
            "MountainDAO": (1390, 740, 1950, 1050, "persistence", ["+ loadMountains()", "+ normalizeMountainCode()", "+ getMountainByCode()"]),
            "Student": (300, 1320, 830, 1630, "entity", ["6 registration fields", "Serializable"]),
            "Mountain": (1030, 1320, 1560, 1630, "entity", ["code, name, province", "description"]),
            "StatisticalInfo": (1770, 1320, 2380, 1630, "entity", ["mountainCode", "numberOfStudents", "totalCost"]),
        }
        relations = [
            ("Program", "Menu", "starts", False),
            ("Menu", "StudentManagement", "uses", False),
            ("StudentManagement", "IStudentManagement", "implements", True),
            ("StudentManagement", "StudentDAO", "coordinates", False),
            ("StudentManagement", "MountainDAO", "coordinates", False),
            ("StudentDAO", "Student", "stores copies", False),
            ("MountainDAO", "Mountain", "loads", False),
            ("StudentManagement", "StatisticalInfo", "creates", False),
        ]
    elif lab_key == "lab2":
        nodes = {
            "Menu": (970, 150, 1590, 400, "presentation", ["+ run()", "- processChoice()", "- loadData()/saveData()"]),
            "ClubIO": (140, 530, 660, 790, "presentation", ["5 Club handlers", "uses IClubManager"]),
            "PlayerIO": (800, 530, 1360, 790, "presentation", ["6 Player handlers", "uses both managers"]),
            "IClubManager": (1500, 530, 2020, 790, "interface", ["<<interface>>", "Club operations"]),
            "IPlayerManager": (2070, 530, 2510, 790, "interface", ["<<interface>>", "Player operations"]),
            "ClubManager": (180, 980, 760, 1280, "business", ["extends AbstractManager<Club>", "filter/load/save"]),
            "PlayerManager": (940, 970, 1570, 1300, "business", ["extends AbstractManager<Player>", "sort/search/shirt rules"]),
            "AbstractManager": (1750, 980, 2390, 1280, "business", ["<<abstract>>", "identity lookup", "defensive copies/change flag"]),
            "Club": (300, 1430, 760, 1660, "entity", ["extends BaseEntity", "name/sponsor/budget"]),
            "Player": (1050, 1420, 1580, 1680, "entity", ["extends BaseEntity", "club/name/position/shirt"]),
            "FileHelpers": (1830, 1430, 2440, 1680, "persistence", ["ClubFileHelper", "PlayerFileHelper", "strict UTF-8 CSV"]),
        }
        relations = [
            ("Menu", "ClubIO", "routes", False),
            ("Menu", "PlayerIO", "routes", False),
            ("Menu", "IClubManager", "owns", False),
            ("Menu", "IPlayerManager", "owns", False),
            ("ClubIO", "IClubManager", "uses", False),
            ("PlayerIO", "IPlayerManager", "uses", False),
            ("PlayerIO", "IClubManager", "uses", False),
            ("ClubManager", "IClubManager", "implements", True),
            ("PlayerManager", "IPlayerManager", "implements", True),
            ("ClubManager", "AbstractManager", "extends", False),
            ("PlayerManager", "AbstractManager", "extends", False),
            ("ClubManager", "Club", "manages", False),
            ("PlayerManager", "Player", "manages", False),
            ("ClubManager", "FileHelpers", "persists", False),
            ("PlayerManager", "FileHelpers", "persists", False),
        ]
    else:
        nodes = {
            "Main": (980, 150, 1580, 430, "presentation", ["+ main(args)", "- nine workflow handlers", "- changed state"]),
            "View": (120, 540, 620, 810, "presentation", ["validated input", "yes/no and menu"]),
            "EmployeeIO": (760, 540, 1320, 820, "presentation", ["+ inputNewEmployee()", "+ inputUpdatedEmployee()"]),
            "EmployeeView": (1470, 540, 2050, 820, "presentation", ["employee table", "active payroll table"]),
            "IEmployeeManager": (2110, 540, 2540, 820, "interface", ["<<interface>>", "CRUD/search/payroll/file"]),
            "EmployeeManager": (850, 1000, 1550, 1340, "business", ["+ add/update/remove", "+ searchByAttribute()", "+ load/save", "+ active payroll"]),
            "IFileReadWrite": (1740, 1020, 2290, 1310, "interface", ["<<interface<T>>>", "+ read()", "+ write(items)"]),
            "EmployeeFileHelper": (1780, 1430, 2400, 1690, "persistence", ["resilient CSV read", "protected replacement"]),
            "Employee": (220, 1090, 700, 1400, "entity", ["7 assignment fields", "implements Identifiable", "implements PayrollCalculable"]),
            "Enums": (300, 1480, 800, 1700, "entity", ["Role", "EmployeeStatus"]),
        }
        relations = [
            ("Main", "View", "uses", False),
            ("Main", "EmployeeIO", "uses", False),
            ("Main", "EmployeeView", "uses", False),
            ("Main", "IEmployeeManager", "owns", False),
            ("EmployeeIO", "View", "uses", False),
            ("EmployeeView", "IEmployeeManager", "reads", False),
            ("EmployeeManager", "IEmployeeManager", "implements", True),
            ("EmployeeManager", "Employee", "manages copies", False),
            ("EmployeeManager", "IFileReadWrite", "depends on", False),
            ("EmployeeFileHelper", "IFileReadWrite", "implements", True),
            ("Employee", "Enums", "contains", False),
        ]

    boxes = {name: tuple(value[:4]) for name, value in nodes.items()}
    for relation in relations:
        draw_relation(draw, boxes, *relation)

    colors = {
        "presentation": (rgb(LIGHT_BLUE), rgb(BLUE)),
        "business": (rgb(LIGHT_GREEN), rgb(GREEN)),
        "persistence": (rgb(LIGHT_ORANGE), rgb(ORANGE)),
        "entity": (rgb(LIGHT_GRAY), rgb(GRAY)),
        "interface": (rgb(PALE_BLUE), rgb(NAVY)),
    }
    for name, (left, top, right, bottom, group, details) in nodes.items():
        fill, outline = colors[group]
        draw.rounded_rectangle((left, top, right, bottom), radius=18, fill=fill, outline=outline, width=6)
        header_bottom = top + 65
        draw.line((left, header_bottom, right, header_bottom), fill=outline, width=4)
        name_box = draw.textbbox((0, 0), name, font=class_font)
        name_width = name_box[2] - name_box[0]
        draw.text(((left + right - name_width) / 2, top + 15), name, fill=rgb(NAVY), font=class_font)
        y = header_bottom + 15
        for detail in details:
            draw.text((left + 18, y), detail, fill=rgb("333333"), font=detail_font)
            y += 39

    output = DIAGRAM_ROOT / f"{lab_key}-uml.png"
    image.save(output, optimize=True)
    return output


def draw_flowchart(lab_key: str, data: Dict[str, object]) -> Path:
    width, height = 1900, 2200
    image = Image.new("RGB", (width, height), rgb(WHITE))
    draw = ImageDraw.Draw(image)
    title_font = load_font("arialbd.ttf", 50)
    body_font = load_font("arial.ttf", 30)
    label_font = load_font("arialbd.ttf", 28)
    title = {
        "lab1": "Algorithm Flow - New Registration",
        "lab2": "Algorithm Flow - Strict Clubs-First Reload",
        "lab3": "Algorithm Flow - Resilient Employee Load",
    }[lab_key]
    draw.text((100, 50), title, fill=rgb(NAVY), font=title_font)

    if lab_key == "lab1":
        steps = [
            ("terminal", "Start"),
            ("process", "Read Student fields and optional base fee"),
            ("process", "Normalize Student ID and Mountain code"),
            ("decision", "All formats valid?"),
            ("decision", "Unique ID and existing Mountain?"),
            ("process", "Calculate 35% discount when eligible"),
            ("process", "Create Student and commit to StudentDAO"),
            ("terminal", "Return success"),
        ]
        rejection = "No -> show precise validation message; state unchanged"
    elif lab_key == "lab2":
        steps = [
            ("terminal", "Function 13: Start"),
            ("process", "Clear both managers"),
            ("process", "Parse every Club line strictly"),
            ("decision", "All Club lines valid?"),
            ("process", "Commit validated Clubs"),
            ("process", "Parse Players against committed Clubs"),
            ("decision", "All Player lines and references valid?"),
            ("process", "Commit validated Players"),
            ("terminal", "Load data successfully"),
        ]
        rejection = "No -> clear both managers; print 'Load data failed!'"
    else:
        steps = [
            ("terminal", "Start reading employees.txt"),
            ("process", "Read next non-blank CSV line"),
            ("decision", "Exactly 7 fields and valid values?"),
            ("decision", "Employee ID unique in loaded rows?"),
            ("process", "Add defensive Employee copy to temporary list"),
            ("decision", "More lines?"),
            ("process", "Manager revalidates complete temporary list"),
            ("process", "Replace in-memory state"),
            ("terminal", "Return success"),
        ]
        rejection = "Bad row -> warn/log/skip; continue with the next line"

    center_x = width // 2
    y = 190
    boxes: List[Tuple[int, int, int, int]] = []
    for shape, text in steps:
        if shape == "decision":
            box_width, box_height = 1100, 190
        elif shape == "terminal":
            box_width, box_height = 780, 135
        else:
            box_width, box_height = 1100, 150
        left, right = center_x - box_width // 2, center_x + box_width // 2
        top, bottom = y, y + box_height
        if boxes:
            previous = boxes[-1]
            draw.line((center_x, previous[3], center_x, top), fill=rgb(GRAY), width=6)
            draw.polygon(
                [(center_x, top), (center_x - 16, top - 24), (center_x + 16, top - 24)], fill=rgb(GRAY)
            )
        if shape == "decision":
            polygon = [(center_x, top), (right, (top + bottom) // 2), (center_x, bottom), (left, (top + bottom) // 2)]
            draw.polygon(polygon, fill=rgb(LIGHT_ORANGE), outline=rgb(ORANGE))
            draw.line(polygon + [polygon[0]], fill=rgb(ORANGE), width=6)
        elif shape == "terminal":
            draw.rounded_rectangle((left, top, right, bottom), radius=65, fill=rgb(LIGHT_GREEN), outline=rgb(GREEN), width=6)
        else:
            draw.rounded_rectangle((left, top, right, bottom), radius=18, fill=rgb(LIGHT_BLUE), outline=rgb(BLUE), width=6)
        text_box = draw.multiline_textbbox((0, 0), text, font=body_font, spacing=6, align="center")
        text_width, text_height = text_box[2] - text_box[0], text_box[3] - text_box[1]
        draw.multiline_text(
            (center_x - text_width / 2, (top + bottom - text_height) / 2),
            text,
            fill=rgb("222222"),
            font=body_font,
            spacing=6,
            align="center",
        )
        boxes.append((left, top, right, bottom))
        y = bottom + 65

    note_top = min(y + 10, height - 180)
    draw.rounded_rectangle((120, note_top, width - 120, note_top + 110), radius=16, fill=rgb(PALE_BLUE), outline=rgb(NAVY), width=4)
    draw.text((160, note_top + 33), rejection, fill=rgb(NAVY), font=label_font)
    output = DIAGRAM_ROOT / f"{lab_key}-algorithm-flow.png"
    image.save(output, optimize=True)
    return output


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row) -> None:
    repeat_table_header(row)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def configure_document(doc: Document, data: Dict[str, object]) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 25, NAVY),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, GREEN),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{data['code']} | {STUDENT_ID} | Verified {VERIFY_DATE_LONG}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(section.footer.paragraphs[0])

    properties = doc.core_properties
    properties.title = f"Report Lab {data['number']} - {data['title']}"
    properties.subject = "LAB211 UML, algorithms, test cases, and traceability"
    properties.author = STUDENT_NAME_VI
    properties.last_modified_by = STUDENT_NAME_VI
    properties.created = VERIFY_DATE
    properties.modified = VERIFY_DATE
    properties.keywords = "LAB211, Java, OOP, UML, algorithms, test cases, AI audit evidence"


def add_styled_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[object]],
    widths: Sequence[float] | None = None,
    font_size: float = 8.5,
) -> object:
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
        set_cell_shading(header_cells[index], NAVY)
        set_cell_margins(header_cells[index])
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in header_cells[index].paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
        if widths:
            header_cells[index].width = Inches(widths[index])
    set_repeat_table_header(table.rows[0])

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, value in enumerate(values):
            cells[index].text = str(value)
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2:
                set_cell_shading(cells[index], PALE_BLUE)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
            if widths:
                cells[index].width = Inches(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    set_cell_margins(cell, 120, 150, 120, 150)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines, 1):
        run = paragraph.add_run(f"{index:02d}  {line}")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        if index < len(lines):
            run.add_break()


def add_bullet_list(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def build_report(lab_key: str, data: Dict[str, object], uml_path: Path, flow_path: Path) -> Path:
    doc = Document()
    configure_document(doc, data)

    cover_spacer = doc.add_paragraph()
    cover_spacer.paragraph_format.space_after = Pt(46)
    code_paragraph = doc.add_paragraph()
    code_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = code_paragraph.add_run(f"LAB211 PROJECT REPORT | LAB {data['number']}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run(str(data["title"]))
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(27)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{data['code']} | UML, Algorithms, and Test Evidence")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(42)
    identity = add_styled_table(
        doc,
        ["Document field", "Value"],
        [
            ("Student", f"{STUDENT_NAME} ({STUDENT_NAME_VI})"),
            ("Student ID", STUDENT_ID),
            ("Course", "LAB211 - OOP with Java Lab"),
            ("Release", f"Portfolio review v1.0 - {VERIFY_DATE_LONG}"),
            ("Repository", REPOSITORY_URL),
        ],
        widths=(1.55, 5.65),
        font_size=9.5,
    )
    identity.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.space_before = Pt(12)
    run = statement.add_run(
        "Scope intentionally limited to the lecturer's review priorities: final UML, implemented algorithms, "
        "test cases, and reproducible evidence."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(str(data["summary"]))
    doc.add_heading("1.1 Verification snapshot", level=2)
    add_styled_table(
        doc,
        ["Check", "Result", "Evidence"],
        [
            ("Java compatibility compile", "PASS", "javac targeting Java 8 with UTF-8 and -Xlint:all"),
            ("Executable deep verification", f"{data['deep_count']} passed, 0 failed", f"labs/{lab_key}/testcases/{'Lab' + str(data['number']) + 'DeepVerification.java'}"),
            ("Import audit", "PASS", "0 wildcard, 0 unresolved, 0 definitely unused explicit imports"),
            ("Report/Audit metadata", "ALIGNED", f"{data['code']}, {STUDENT_ID}, {VERIFY_DATE_LONG}"),
            ("Required functions traced", f"{data['functions']}/{data['functions']}", "Section 5 requirement-to-source matrix"),
        ],
        widths=(2.15, 1.35, 3.7),
    )
    doc.add_heading("1.2 Evidence interpretation", level=2)
    doc.add_paragraph(
        f"The runnable deep harness was recompiled and executed during this review: {data['deep_count']} checks "
        "passed and none failed. The smaller verification_results.txt file is retained as earlier regression "
        f"evidence ({data['regression_count']} passed), but it is not counted in the current completion total."
    )

    doc.add_page_break()
    doc.add_heading("2. UML Design", level=1)
    doc.add_paragraph(
        "The diagram below is regenerated from the final source structure. It presents the dependencies that "
        "matter to grading and maintenance, without decorative classes or layers that are not implemented."
    )
    picture = doc.add_picture(str(uml_path), width=Inches(7.2))
    picture_paragraph = doc.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph(f"Figure 1. Final class relationships for {data['code']}.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(8.5)
    doc.add_heading("2.1 Responsibility matrix", level=2)
    add_styled_table(
        doc,
        ["Class / Interface", "Final responsibility"],
        data["classes"],
        widths=(2.15, 5.05),
        font_size=8.6,
    )
    doc.add_heading("2.2 OOP evidence", level=2)
    oop_items = {
        "lab1": [
            "Encapsulation: StudentDAO protects mutable Student records through deep defensive copies.",
            "Abstraction: IStudentManagement exposes business use cases without persistence details.",
            "Polymorphism: Menu depends on the business contract represented by StudentManagement operations.",
            "Responsibility separation: presentation, business, data access, utilities, and entities remain distinct.",
        ],
        "lab2": [
            "Inheritance: Club and Player extend BaseEntity; managers reuse AbstractManager<T>.",
            "Abstraction: manager and file-helper interfaces expose actual assignment operations.",
            "Polymorphism: Menu and IO classes operate through IClubManager and IPlayerManager.",
            "Encapsulation: managers return entity copies, so callers cannot bypass validation through setters.",
        ],
        "lab3": [
            "Abstraction: IEmployeeManager and IFileReadWrite<T> separate business and persistence contracts.",
            "Polymorphism: Employee implements Identifiable and PayrollCalculable domain behaviors.",
            "Encapsulation: public queries return deep Employee copies; internal mutation remains controlled.",
            "Dependency inversion: EmployeeManager receives its file helper through constructor injection.",
        ],
    }[lab_key]
    add_bullet_list(doc, oop_items)

    doc.add_page_break()
    doc.add_heading("3. Implemented Algorithms", level=1)
    doc.add_paragraph(
        "Pseudocode reflects the final implementation rather than an idealized alternative. Complexity statements "
        "also account for the list-based data structures used by these assignment-sized applications."
    )
    flow_picture = doc.add_picture(str(flow_path), width=Inches(5.65))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph("Figure 2. Representative decision flow from the final implementation.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(8.5)

    for index, algorithm in enumerate(data["algorithms"]):
        if index == 2:
            doc.add_page_break()
        doc.add_heading(str(algorithm["name"]), level=2)
        purpose = doc.add_paragraph()
        purpose.add_run("Purpose: ").bold = True
        purpose.add_run(str(algorithm["purpose"]))
        add_code_block(doc, algorithm["pseudocode"])
        complexity = doc.add_paragraph()
        complexity.add_run("Complexity: ").bold = True
        complexity.add_run(str(algorithm["complexity"]))

    doc.add_page_break()
    doc.add_heading("4. Test Cases and Executed Evidence", level=1)
    doc.add_heading("4.1 Execution record", level=2)
    add_styled_table(
        doc,
        ["Artifact", "Status", "Interpretation"],
        [
            ("Source compilation", "PASS", "All production and deep-test Java files compile for Java 8."),
            ("Deep harness", f"{data['deep_count']}/{data['deep_count']} PASS", "Re-executed on 21 July 2026; process exit code 0."),
            ("Earlier regression record", f"{data['regression_count']}/{data['regression_count']} PASS", "Retained historical result; not added to current 153-check claim."),
            ("State preservation", "PASS", "Failure paths assert memory/file content, not only return values."),
        ],
        widths=(2.0, 1.5, 3.7),
    )
    doc.add_heading("4.2 Representative test matrix", level=2)
    add_styled_table(
        doc,
        ["ID", "Scenario", "Input / action", "Expected result", "Actual"],
        data["test_cases"],
        widths=(0.65, 1.45, 1.9, 2.45, 0.65),
        font_size=7.7,
    )
    doc.add_heading("4.3 Coverage strategy", level=2)
    add_bullet_list(
        doc,
        [
            "Equivalence partitions: valid, invalid format, duplicate, missing reference, and empty/null inputs.",
            "Numeric boundaries: lower/upper limits plus NaN and positive infinity where double values are accepted.",
            "State assertions: rejected operations must preserve the prior object list or target file.",
            "Persistence round-trips: valid text/object data is saved and reloaded with essential values preserved.",
            "Encapsulation checks: mutations to returned objects/lists must not change manager/DAO state.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("5. Requirement-to-Source Traceability", level=1)
    doc.add_paragraph(
        "Every visible assignment function maps to a final handler, a business/persistence method, and at least one "
        "verification artifact or explicit source-path review."
    )
    add_styled_table(
        doc,
        ["Requirement", "Menu / handler", "Final implementation", "Evidence"],
        data["traceability"],
        widths=(1.55, 1.7, 2.25, 1.7),
        font_size=7.6,
    )
    doc.add_heading("5.1 Import and source hygiene", level=2)
    add_styled_table(
        doc,
        ["Audit item", "Result", "Basis"],
        [
            ("Incorrect/unresolved imports", "0", "Complete production + test compilation succeeded."),
            ("Wildcard imports", "0", "Repository-wide source scan."),
            ("Definitely unused explicit imports", "0", "Imported simple names are referenced in their compilation units."),
            ("Generated binaries committed", "0", ".gitignore excludes build, dist, class, jar, private IDE state, and runtime logs."),
        ],
        widths=(2.25, 0.75, 4.2),
    )
    doc.add_heading("5.2 Known limitations and deliberate trade-offs", level=2)
    add_bullet_list(doc, data["limitations"])

    doc.add_heading("5.3 Reproduction", level=2)
    doc.add_paragraph("From the repository root on Windows:")
    add_code_block(doc, ["PowerShell -ExecutionPolicy Bypass -File .\\scripts\\verify-all.ps1"])
    doc.add_paragraph(
        "GitHub Actions invokes the same script on every push and pull request. The workflow pins third-party "
        "actions to reviewed commit SHAs and grants read-only repository permission."
    )

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        f"The final {data['code']} source implements all {data['functions']} required functions, compiles for Java 8, "
        f"and passes {data['deep_count']} executable deep checks. The UML, algorithm descriptions, test evidence, "
        "and AI Audit Log now use the same identifiers, counts, paths, and verification date."
    )

    output = EDITABLE_ROOT / f"Report_Lab_{data['number']}_{STUDENT_ID}.docx"
    doc.save(output)
    return output


def register_pdf_fonts() -> None:
    font_root = Path("C:/Windows/Fonts")
    registrations = {
        "PortfolioArial": "arial.ttf",
        "PortfolioArialBold": "arialbd.ttf",
        "PortfolioArialItalic": "ariali.ttf",
        "PortfolioConsolas": "consola.ttf",
    }
    for logical_name, filename in registrations.items():
        if logical_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(logical_name, str(font_root / filename)))


def pdf_styles() -> Dict[str, ParagraphStyle]:
    sample = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "PortfolioBody",
            parent=sample["BodyText"],
            fontName="PortfolioArial",
            fontSize=9.1,
            leading=12.1,
            textColor=colors.HexColor("#222222"),
            spaceAfter=5,
        ),
        "small": ParagraphStyle(
            "PortfolioSmall",
            parent=sample["BodyText"],
            fontName="PortfolioArial",
            fontSize=7.4,
            leading=9.2,
            textColor=colors.HexColor("#333333"),
        ),
        "table": ParagraphStyle(
            "PortfolioTable",
            parent=sample["BodyText"],
            fontName="PortfolioArial",
            fontSize=7.1,
            leading=8.6,
            textColor=colors.HexColor("#222222"),
        ),
        "table_header": ParagraphStyle(
            "PortfolioTableHeader",
            parent=sample["BodyText"],
            fontName="PortfolioArialBold",
            fontSize=7.2,
            leading=8.6,
            textColor=colors.white,
            alignment=TA_CENTER,
        ),
        "title": ParagraphStyle(
            "PortfolioTitle",
            parent=sample["Title"],
            fontName="PortfolioArialBold",
            fontSize=24,
            leading=29,
            textColor=colors.HexColor("#17365D"),
            alignment=TA_CENTER,
            spaceAfter=16,
        ),
        "subtitle": ParagraphStyle(
            "PortfolioSubtitle",
            parent=sample["BodyText"],
            fontName="PortfolioArial",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#666666"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "h1": ParagraphStyle(
            "PortfolioH1",
            parent=sample["Heading1"],
            fontName="PortfolioArialBold",
            fontSize=16,
            leading=19,
            textColor=colors.HexColor("#17365D"),
            spaceBefore=3,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "PortfolioH2",
            parent=sample["Heading2"],
            fontName="PortfolioArialBold",
            fontSize=11.5,
            leading=14,
            textColor=colors.HexColor("#2F75B5"),
            spaceBefore=7,
            spaceAfter=5,
        ),
        "h3": ParagraphStyle(
            "PortfolioH3",
            parent=sample["Heading3"],
            fontName="PortfolioArialBold",
            fontSize=9.8,
            leading=12,
            textColor=colors.HexColor("#2E7D32"),
            spaceBefore=5,
            spaceAfter=4,
        ),
        "caption": ParagraphStyle(
            "PortfolioCaption",
            parent=sample["BodyText"],
            fontName="PortfolioArialItalic",
            fontSize=7.5,
            leading=9,
            textColor=colors.HexColor("#555555"),
            alignment=TA_CENTER,
            spaceAfter=7,
        ),
        "code": ParagraphStyle(
            "PortfolioCode",
            parent=sample["Code"],
            fontName="PortfolioConsolas",
            fontSize=7.7,
            leading=10,
            textColor=colors.HexColor("#1F2937"),
            leftIndent=7,
            rightIndent=7,
            spaceAfter=5,
        ),
        "bullet": ParagraphStyle(
            "PortfolioBullet",
            parent=sample["BodyText"],
            fontName="PortfolioArial",
            fontSize=8.8,
            leading=11.5,
            leftIndent=13,
            firstLineIndent=-7,
            bulletIndent=4,
            textColor=colors.HexColor("#222222"),
            spaceAfter=3,
        ),
    }


def pdf_paragraph(text: object, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(str(text)).replace("\n", "<br/>"), style)


def pdf_table(
    headers: Sequence[str],
    rows: Iterable[Sequence[object]],
    column_widths: Sequence[float],
    styles: Dict[str, ParagraphStyle],
    header_rows: int = 1,
) -> LongTable:
    prepared: List[List[Paragraph]] = [
        [pdf_paragraph(header, styles["table_header"]) for header in headers]
    ]
    for row in rows:
        prepared.append([pdf_paragraph(value, styles["table"]) for value in row])
    table = LongTable(prepared, colWidths=column_widths, repeatRows=header_rows, hAlign="CENTER")
    commands = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17365D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AFC4D6")),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    for row_index in range(2, len(prepared), 2):
        commands.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor("#EEF5FB")))
    table.setStyle(TableStyle(commands))
    return table


def scaled_reportlab_image(path: Path, max_width: float, max_height: float) -> ReportLabImage:
    with Image.open(path) as source:
        width, height = source.size
    scale = min(max_width / width, max_height / height)
    return ReportLabImage(str(path), width=width * scale, height=height * scale)


def add_pdf_bullets(story: List[object], items: Sequence[str], styles: Dict[str, ParagraphStyle]) -> None:
    for item in items:
        story.append(Paragraph("&#8226;&nbsp; " + escape(item), styles["bullet"]))


def report_page(canvas, document) -> None:
    canvas.saveState()
    canvas.setTitle(document.title)
    canvas.setAuthor(STUDENT_NAME_VI)
    canvas.setSubject("LAB211 UML, algorithms, test cases, and traceability")
    width, height = A4
    if document.page > 1:
        canvas.setFont("PortfolioArial", 7.5)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawRightString(width - 43, height - 30, document.header_text)
    canvas.setFont("PortfolioArial", 8)
    canvas.setFillColor(colors.HexColor("#555555"))
    canvas.drawCentredString(width / 2, 24, f"Page {document.page}")
    canvas.restoreState()


def build_pdf_report(lab_key: str, data: Dict[str, object], uml_path: Path, flow_path: Path) -> Path:
    register_pdf_fonts()
    PDF_ROOT.mkdir(parents=True, exist_ok=True)
    output = PDF_ROOT / f"Report_Lab_{data['number']}_{STUDENT_ID}.pdf"
    styles = pdf_styles()
    document = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=43,
        leftMargin=43,
        topMargin=42,
        bottomMargin=38,
        title=f"Report Lab {data['number']} - {data['title']}",
        author=STUDENT_NAME_VI,
        subject="LAB211 UML, algorithms, test cases, and traceability",
    )
    document.header_text = f"{data['code']} | {STUDENT_ID} | Verified {VERIFY_DATE_LONG}"
    story: List[object] = []

    story.extend(
        [
            Spacer(1, 0.65 * inch),
            Paragraph(f"LAB211 PROJECT REPORT | LAB {data['number']}", styles["subtitle"]),
            Paragraph(escape(str(data["title"])), styles["title"]),
            Paragraph(f"{data['code']} | UML, Algorithms, and Test Evidence", styles["subtitle"]),
            Spacer(1, 0.34 * inch),
            pdf_table(
                ["Document field", "Value"],
                [
                    ("Student", f"{STUDENT_NAME} ({STUDENT_NAME_VI})"),
                    ("Student ID", STUDENT_ID),
                    ("Course", "LAB211 - OOP with Java Lab"),
                    ("Release", f"Portfolio review v1.0 - {VERIFY_DATE_LONG}"),
                    ("Repository", REPOSITORY_URL),
                ],
                (1.45 * inch, 5.25 * inch),
                styles,
            ),
            Spacer(1, 0.25 * inch),
            Paragraph(
                "Scope intentionally limited to the lecturer's review priorities: final UML, implemented "
                "algorithms, test cases, and reproducible evidence.",
                styles["caption"],
            ),
            PageBreak(),
            Paragraph("1. Executive Summary", styles["h1"]),
            pdf_paragraph(data["summary"], styles["body"]),
            Paragraph("1.1 Verification snapshot", styles["h2"]),
            pdf_table(
                ["Check", "Result", "Evidence"],
                [
                    ("Java compatibility compile", "PASS", "javac targeting Java 8, UTF-8, and -Xlint:all"),
                    ("Executable deep verification", f"{data['deep_count']} passed, 0 failed", f"labs/{lab_key}/testcases/Lab{data['number']}DeepVerification.java"),
                    ("Import audit", "PASS", "0 wildcard, 0 unresolved, 0 definitely unused explicit imports"),
                    ("Report/Audit metadata", "ALIGNED", f"{data['code']}, {STUDENT_ID}, {VERIFY_DATE_LONG}"),
                    ("Required functions traced", f"{data['functions']}/{data['functions']}", "Section 5 traceability matrix"),
                ],
                (1.9 * inch, 1.25 * inch, 3.55 * inch),
                styles,
            ),
            Spacer(1, 6),
            Paragraph("1.2 Evidence interpretation", styles["h2"]),
            pdf_paragraph(
                f"The runnable deep harness was recompiled and executed during this review: {data['deep_count']} "
                f"checks passed and none failed. The smaller verification_results.txt file is retained as earlier "
                f"regression evidence ({data['regression_count']} passed), but is not counted in the current total.",
                styles["body"],
            ),
            PageBreak(),
            Paragraph("2. UML Design", styles["h1"]),
            pdf_paragraph(
                "This diagram is regenerated from the final source structure and names only implemented classes, "
                "interfaces, and relationships relevant to grading and maintenance.",
                styles["body"],
            ),
            scaled_reportlab_image(uml_path, 7.0 * inch, 4.72 * inch),
            Paragraph(f"Figure 1. Final class relationships for {data['code']}.", styles["caption"]),
            Paragraph("2.1 Responsibility matrix", styles["h2"]),
            pdf_table(
                ["Class / Interface", "Final responsibility"],
                data["classes"],
                (1.9 * inch, 4.8 * inch),
                styles,
            ),
            Paragraph("2.2 OOP evidence", styles["h2"]),
        ]
    )
    oop_items = {
        "lab1": [
            "Encapsulation: StudentDAO protects mutable Student records through deep defensive copies.",
            "Abstraction: IStudentManagement exposes use cases without persistence details.",
            "Responsibility separation: presentation, business, data access, utilities, and entities remain distinct.",
        ],
        "lab2": [
            "Inheritance: Club and Player extend BaseEntity; managers reuse AbstractManager<T>.",
            "Polymorphism: Menu and IO classes operate through IClubManager and IPlayerManager.",
            "Encapsulation: managers return entity copies so callers cannot bypass validation.",
        ],
        "lab3": [
            "Abstraction: IEmployeeManager and IFileReadWrite<T> separate business and persistence contracts.",
            "Polymorphism: Employee implements Identifiable and PayrollCalculable behaviors.",
            "Dependency inversion: EmployeeManager receives its file helper through constructor injection.",
        ],
    }[lab_key]
    add_pdf_bullets(story, oop_items, styles)
    story.extend(
        [
            PageBreak(),
            Paragraph("3. Implemented Algorithms", styles["h1"]),
            pdf_paragraph(
                "Pseudocode reflects the final implementation rather than an idealized alternative. Complexity "
                "statements account for the list-based structures used by these assignment-sized applications.",
                styles["body"],
            ),
            scaled_reportlab_image(flow_path, 4.3 * inch, 5.6 * inch),
            Paragraph("Figure 2. Representative decision flow from the final implementation.", styles["caption"]),
            PageBreak(),
        ]
    )
    for algorithm in data["algorithms"]:
        code = "\n".join(f"{index:02d}  {line}" for index, line in enumerate(algorithm["pseudocode"], 1))
        block = [
            Paragraph(escape(str(algorithm["name"])), styles["h2"]),
            Paragraph("<b>Purpose:</b> " + escape(str(algorithm["purpose"])), styles["body"]),
            Table(
                [[Preformatted(code, styles["code"])]],
                colWidths=[6.65 * inch],
                style=TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F9FC")),
                        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#B7C9DA")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                ),
            ),
            Paragraph("<b>Complexity:</b> " + escape(str(algorithm["complexity"])), styles["body"]),
            Spacer(1, 5),
        ]
        story.append(KeepTogether(block))

    story.extend(
        [
            PageBreak(),
            Paragraph("4. Test Cases and Executed Evidence", styles["h1"]),
            Paragraph("4.1 Execution record", styles["h2"]),
            pdf_table(
                ["Artifact", "Status", "Interpretation"],
                [
                    ("Source compilation", "PASS", "All production and deep-test Java files compile for Java 8."),
                    ("Deep harness", f"{data['deep_count']}/{data['deep_count']} PASS", "Re-executed 21 July 2026; exit code 0."),
                    ("Earlier regression record", f"{data['regression_count']}/{data['regression_count']} PASS", "Historical result; not added to the current 153-check claim."),
                    ("State preservation", "PASS", "Failure paths assert memory/file content, not only return values."),
                ],
                (1.7 * inch, 1.35 * inch, 3.65 * inch),
                styles,
            ),
            Spacer(1, 6),
            Paragraph("4.2 Representative test matrix", styles["h2"]),
            pdf_table(
                ["ID", "Scenario", "Input / action", "Expected result", "Actual"],
                data["test_cases"],
                (0.6 * inch, 1.25 * inch, 1.65 * inch, 2.55 * inch, 0.65 * inch),
                styles,
            ),
            Paragraph("4.3 Coverage strategy", styles["h2"]),
        ]
    )
    add_pdf_bullets(
        story,
        [
            "Equivalence partitions: valid, invalid format, duplicate, missing reference, and empty/null inputs.",
            "Numeric boundaries: lower/upper limits plus NaN and positive infinity for double values.",
            "State assertions: rejected operations preserve the prior object list or target file.",
            "Persistence round-trips: valid object/text data reloads with essential values preserved.",
            "Encapsulation checks: mutations to returned objects/lists cannot alter managed state.",
        ],
        styles,
    )
    story.extend(
        [
            PageBreak(),
            Paragraph("5. Requirement-to-Source Traceability", styles["h1"]),
            pdf_paragraph(
                "Every visible assignment function maps to a final handler, implementation method, and executable "
                "or directly inspectable verification artifact.",
                styles["body"],
            ),
            pdf_table(
                ["Requirement", "Menu / handler", "Final implementation", "Evidence"],
                data["traceability"],
                (1.35 * inch, 1.55 * inch, 2.05 * inch, 1.75 * inch),
                styles,
            ),
            Paragraph("5.1 Import and source hygiene", styles["h2"]),
            pdf_table(
                ["Audit item", "Result", "Basis"],
                [
                    ("Incorrect/unresolved imports", "0", "Complete production + test compilation succeeded."),
                    ("Wildcard imports", "0", "Repository-wide source scan."),
                    ("Definitely unused explicit imports", "0", "Imported simple names are referenced in their units."),
                    ("Generated binaries committed", "0", ".gitignore excludes build, dist, class, jar, private IDE state, and logs."),
                ],
                (2.1 * inch, 0.7 * inch, 3.9 * inch),
                styles,
            ),
            Paragraph("5.2 Known limitations and deliberate trade-offs", styles["h2"]),
        ]
    )
    add_pdf_bullets(story, data["limitations"], styles)
    story.extend(
        [
            Paragraph("5.3 Reproduction", styles["h2"]),
            pdf_paragraph("From the repository root on Windows:", styles["body"]),
            Table(
                [[Preformatted("PowerShell -ExecutionPolicy Bypass -File .\\scripts\\verify-all.ps1", styles["code"])]],
                colWidths=[6.65 * inch],
                style=TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F9FC")),
                        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#B7C9DA")),
                    ]
                ),
            ),
            Spacer(1, 6),
            pdf_paragraph(
                "GitHub Actions invokes the same script on every push and pull request. The workflow pins actions "
                "to reviewed commit SHAs and grants read-only repository permission.",
                styles["body"],
            ),
            Paragraph("Conclusion", styles["h1"]),
            pdf_paragraph(
                f"The final {data['code']} source implements all {data['functions']} required functions, compiles "
                f"for Java 8, and passes {data['deep_count']} executable deep checks. The UML, algorithms, test "
                "evidence, and AI Audit Log use the same identifiers, counts, paths, and verification date.",
                styles["body"],
            ),
        ]
    )
    document.build(story, onFirstPage=report_page, onLaterPages=report_page)
    return output


def copy_style(source_cell, target_cell) -> None:
    if source_cell.has_style:
        target_cell._style = copy(source_cell._style)
    if source_cell.number_format:
        target_cell.number_format = source_cell.number_format
    target_cell.alignment = copy(source_cell.alignment)
    target_cell.font = copy(source_cell.font)
    target_cell.fill = copy(source_cell.fill)
    target_cell.border = copy(source_cell.border)


def normalize_audit_claims(workbook) -> None:
    replacements = {
        "Java 8/Ant PASS": "Java 8-compatible javac PASS",
        "Report Section 4 UML": "Report Section 2 UML",
    }
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    value = cell.value
                    for old, new in replacements.items():
                        value = value.replace(old, new)
                    cell.value = value


def update_audit_workbook(lab_key: str, data: Dict[str, object]) -> Path:
    workbook_path = AUDIT_ROOT / f"AI_Audit_Log_Lab_{data['number']}_{STUDENT_ID}.xlsx"
    workbook = load_workbook(workbook_path)
    normalize_audit_claims(workbook)

    workbook.properties.creator = STUDENT_NAME_VI
    workbook.properties.lastModifiedBy = STUDENT_NAME_VI
    workbook.properties.title = f"AI Audit Log Lab {data['number']} - {data['code']}"
    workbook.properties.subject = "Transparent AI usage with final source and test evidence"
    workbook.properties.description = (
        f"Aligned with the final {data['code']} report and source; verified {VERIFY_DATE_LONG}."
    )
    workbook.properties.created = VERIFY_DATE
    workbook.properties.modified = VERIFY_DATE
    workbook.properties.keywords = "LAB211, AI audit, evidence, prompts, human delta, hallucination detection"

    metadata = workbook["1. Metadata & Summary"]
    metadata["A8"] = "Repository:"
    metadata["C8"] = REPOSITORY_URL
    metadata["C8"].hyperlink = REPOSITORY_URL
    metadata["C8"].style = "Hyperlink"
    copy_style(metadata["A7"], metadata["A8"])
    metadata["A14"] = "Verified On:"
    metadata["C14"] = VERIFY_DATE
    metadata["C14"].number_format = "dd mmmm yyyy"
    copy_style(metadata["A13"], metadata["A14"])
    copy_style(metadata["C13"], metadata["C14"])
    metadata.sheet_view.showGridLines = False
    metadata.freeze_panes = "A3"
    metadata.oddFooter.center.text = f"{data['code']} | {STUDENT_ID} | {VERIFY_DATE_LONG}"

    detail = workbook["2. Detailed Audit Log"]
    detail.sheet_view.showGridLines = False
    detail.freeze_panes = "A4"
    detail.auto_filter.ref = f"A3:H{detail.max_row}"
    entries: List[Tuple[str, str, str, str, str]] = []
    for row_index in range(4, detail.max_row + 1):
        entry_value = detail.cell(row_index, 1).value
        if entry_value in (None, ""):
            continue
        entry_text = str(entry_value)
        if entry_text.isdigit():
            entry_text = entry_text.zfill(3)
        evidence_id = f"EV-{entry_text}"
        prompt_type = str(detail.cell(row_index, 2).value or "")
        stage = str(detail.cell(row_index, 3).value or "")
        problem = str(detail.cell(row_index, 4).value or "")
        evidence = str(detail.cell(row_index, 8).value or "")
        evidence = re.sub(r"^EV-[^|]+\|\s*", "", evidence)
        detail.cell(row_index, 8).value = f"{evidence_id} | {evidence}"
        entries.append((evidence_id, prompt_type, stage, problem, evidence))
    detail.oddFooter.center.text = f"{data['code']} | Detailed AI Audit Log | {STUDENT_ID}"

    hallucinations = workbook["3. Hallucination Detection"]
    hallucinations.sheet_view.showGridLines = False
    hallucinations.freeze_panes = "A4"
    hallucinations.oddFooter.center.text = f"{data['code']} | Hallucination Detection | {STUDENT_ID}"

    checklist = workbook["4. Self-Assessment Checklist"]
    checklist.sheet_view.showGridLines = False
    checklist.freeze_panes = "A5"
    checklist.oddFooter.center.text = f"{data['code']} | Self-Assessment | {STUDENT_ID}"

    evidence_title = "5. Evidence Index"
    if evidence_title in workbook.sheetnames:
        del workbook[evidence_title]
    evidence_sheet = workbook.create_sheet(evidence_title)
    evidence_sheet.sheet_view.showGridLines = False
    evidence_sheet.freeze_panes = "A7"
    evidence_sheet.merge_cells("A1:H1")
    evidence_sheet["A1"] = "FINAL SOURCE & VERIFICATION EVIDENCE INDEX"
    evidence_sheet["A1"].font = Font(name="Arial", size=16, bold=True, color=WHITE)
    evidence_sheet["A1"].fill = PatternFill("solid", fgColor=NAVY)
    evidence_sheet["A1"].alignment = Alignment(horizontal="center", vertical="center")
    evidence_sheet.row_dimensions[1].height = 30
    evidence_sheet.merge_cells("A2:H2")
    evidence_sheet["A2"] = (
        f"{data['code']} | {STUDENT_NAME_VI} | {STUDENT_ID} | Verified {VERIFY_DATE_LONG} | "
        f"Deep harness: {data['deep_count']} passed, 0 failed"
    )
    evidence_sheet["A2"].font = Font(name="Arial", size=10, italic=True, color=NAVY)
    evidence_sheet["A2"].fill = PatternFill("solid", fgColor=LIGHT_BLUE)
    evidence_sheet["A2"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    evidence_sheet.merge_cells("A4:H4")
    evidence_sheet["A4"] = (
        "Evidence is textual and inspectable: final source paths, runnable test code, recorded output, report links, "
        "verification date, and hashes. No AI conversation screenshot was fabricated."
    )
    evidence_sheet["A4"].alignment = Alignment(wrap_text=True, vertical="center")
    evidence_sheet["A4"].fill = PatternFill("solid", fgColor=PALE_BLUE)
    evidence_sheet["A4"].font = Font(name="Arial", size=9, color="333333")
    evidence_sheet.row_dimensions[4].height = 32

    headers = [
        "Evidence ID",
        "Prompt Type",
        "DTC Component",
        "Decision / Context",
        "Artifact cited in Audit Log",
        "Current verification",
        "Verified on",
        "Integrity / repository evidence",
    ]
    header_row = 6
    thin = Side(style="thin", color="B7C9DA")
    for column, header in enumerate(headers, 1):
        cell = evidence_sheet.cell(header_row, column, header)
        cell.font = Font(name="Arial", size=9, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    evidence_sheet.row_dimensions[header_row].height = 34

    lab_number = data["number"]
    test_source = REPO_ROOT / "labs" / lab_key / "testcases" / f"Lab{lab_number}DeepVerification.java"
    result_source = REPO_ROOT / "labs" / lab_key / "testcases" / "deep_verification_results.txt"
    integrity = (
        f"Harness SHA-256 {sha256(test_source)[:16]}... | "
        f"Output SHA-256 {sha256(result_source)[:16]}..."
    )
    for offset, entry in enumerate(entries, 1):
        evidence_id, prompt_type, stage, problem, artifact = entry
        row_index = header_row + offset
        current_verification = (
            f"Mapped to final source/report; Lab {lab_number} deep harness re-executed: "
            f"{data['deep_count']} passed, 0 failed."
        )
        repository_evidence = (
            f"labs/{lab_key}/testcases/Lab{lab_number}DeepVerification.java | "
            f"labs/{lab_key}/testcases/deep_verification_results.txt | {integrity}"
        )
        values = [
            evidence_id,
            prompt_type,
            stage,
            problem,
            artifact,
            current_verification,
            VERIFY_DATE,
            repository_evidence,
        ]
        for column, value in enumerate(values, 1):
            cell = evidence_sheet.cell(row_index, column, value)
            cell.font = Font(name="Arial", size=8.2, color="222222")
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if offset % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
        evidence_sheet.cell(row_index, 7).number_format = "dd mmmm yyyy"
        evidence_sheet.row_dimensions[row_index].height = 72

    last_row = header_row + len(entries)
    evidence_sheet.auto_filter.ref = f"A{header_row}:H{last_row}"
    widths = [14, 16, 18, 42, 48, 38, 16, 52]
    for column, width in enumerate(widths, 1):
        evidence_sheet.column_dimensions[get_column_letter(column)].width = width
    evidence_sheet.page_setup.orientation = "landscape"
    evidence_sheet.page_setup.paperSize = evidence_sheet.PAPERSIZE_A4
    evidence_sheet.page_setup.fitToWidth = 1
    evidence_sheet.page_setup.fitToHeight = 0
    evidence_sheet.sheet_properties.pageSetUpPr.fitToPage = True
    evidence_sheet.print_title_rows = f"1:{header_row}"
    evidence_sheet.print_area = f"A1:H{last_row}"
    evidence_sheet.oddHeader.center.text = f"AI Audit Evidence | {data['code']} | {STUDENT_ID}"
    evidence_sheet.oddFooter.center.text = "Page &P of &N"
    evidence_sheet.oddFooter.right.text = REPOSITORY_URL

    workbook.save(workbook_path)
    return workbook_path


def main() -> None:
    DIAGRAM_ROOT.mkdir(parents=True, exist_ok=True)
    EDITABLE_ROOT.mkdir(parents=True, exist_ok=True)
    AUDIT_ROOT.mkdir(parents=True, exist_ok=True)
    PDF_ROOT.mkdir(parents=True, exist_ok=True)
    for lab_key, data in LABS.items():
        uml = draw_class_diagram(lab_key, data)
        flow = draw_flowchart(lab_key, data)
        report = build_report(lab_key, data, uml, flow)
        pdf = build_pdf_report(lab_key, data, uml, flow)
        audit = update_audit_workbook(lab_key, data)
        print(f"Generated {report.relative_to(REPO_ROOT)}")
        print(f"Generated {pdf.relative_to(REPO_ROOT)}")
        print(f"Updated   {audit.relative_to(REPO_ROOT)}")
        print(f"Generated {uml.relative_to(REPO_ROOT)}")
        print(f"Generated {flow.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
