#!/usr/bin/env python3
"""Validate report, Audit Log, and source alignment for the portfolio release."""

from __future__ import annotations

import hashlib
import re
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
SHEETS = [
    "1. Metadata & Summary",
    "2. Detailed Audit Log",
    "3. Hallucination Detection",
    "4. Self-Assessment Checklist",
    "5. Evidence Index",
]
LABS = {
    1: {
        "code": "J1.L.P0027",
        "assignment": "J1.L.P0027 - Mountain Hiking Challenge Registration",
        "passed": 34,
        "prompts": 12,
    },
    2: {
        "code": "J1.L.P0036",
        "assignment": "J1.L.P0036 - Football Club & Player Management",
        "passed": 62,
        "prompts": 18,
    },
    3: {
        "code": "J1.L.P0037",
        "assignment": "J1.L.P0037 - Employee Payroll Management System",
        "passed": 57,
        "prompts": 20,
    },
}


def require(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def workbook_values(workbook) -> list[str]:
    values: list[str] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            values.extend(str(cell.value) for cell in row if cell.value is not None)
    return values


def metadata_value(sheet, label: str):
    for row in sheet.iter_rows():
        if row[0].value == label:
            return row[2].value
    raise AssertionError(f"Missing Audit Log metadata label: {label}")


def validate_java_imports() -> int:
    java_files = sorted((ROOT / "labs").glob("lab*/**/*.java"))
    require(bool(java_files), "No Java files found.")
    for path in java_files:
        text = path.read_text(encoding="utf-8")
        require(
            not re.search(r"^\s*import\s+(?:static\s+)?[\w.]+\.\*\s*;", text, re.MULTILINE),
            f"Wildcard import found: {path.relative_to(ROOT)}",
        )
        body = re.sub(r"^\s*import\s+[^;]+;\s*$", "", text, flags=re.MULTILINE)
        for match in re.finditer(r"^\s*import\s+(?:static\s+)?([\w.]+)\s*;", text, re.MULTILINE):
            simple_name = match.group(1).rsplit(".", 1)[-1]
            require(
                re.search(rf"\b{re.escape(simple_name)}\b", body) is not None,
                f"Definitely unused explicit import {simple_name}: {path.relative_to(ROOT)}",
            )
    return len(java_files)


def validate_release_checksums() -> int:
    checksum_path = ROOT / "evidence" / "SHA256SUMS.txt"
    require(checksum_path.is_file(), "Missing evidence/SHA256SUMS.txt.")
    entries = 0
    for line in checksum_path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        expected_hash, relative_path = line.split(maxsplit=1)
        artifact_path = ROOT / relative_path
        require(artifact_path.is_file(), f"Checksum target missing: {relative_path}")
        require(sha256(artifact_path) == expected_hash, f"Checksum mismatch: {relative_path}")
        entries += 1
    require(entries == 9, "Release checksum list must contain nine report/Audit artifacts.")
    return entries


def validate_lab(lab_number: int, expected: dict[str, object]) -> None:
    lab = f"lab{lab_number}"
    result = f"{expected['passed']} passed, 0 failed"
    pdf_path = ROOT / "output" / "pdf" / f"Report_Lab_{lab_number}_SE204330.pdf"
    docx_path = ROOT / "docs" / "editable" / f"Report_Lab_{lab_number}_SE204330.docx"
    audit_path = ROOT / "docs" / "ai-audit-logs" / f"AI_Audit_Log_Lab_{lab_number}_SE204330.xlsx"
    harness_path = ROOT / "labs" / lab / "testcases" / f"Lab{lab_number}DeepVerification.java"
    output_path = ROOT / "labs" / lab / "testcases" / "deep_verification_results.txt"

    for path in (pdf_path, docx_path, audit_path, harness_path, output_path):
        require(path.is_file(), f"Missing required artifact: {path.relative_to(ROOT)}")

    reader = PdfReader(pdf_path)
    require(len(reader.pages) == 7, f"Lab {lab_number} PDF must have 7 pages.")
    pdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    for token in ("SE204330", str(expected["code"]), result, "UML", "ALGORITHM", "TEST"):
        require(token.lower() in pdf_text.lower(), f"Lab {lab_number} PDF is missing: {token}")

    doc = Document(docx_path)
    docx_text = document_text(doc)
    for token in ("SE204330", str(expected["code"]), result, "UML", "Algorithm", "Test"):
        require(token.lower() in docx_text.lower(), f"Lab {lab_number} DOCX is missing: {token}")

    workbook = load_workbook(audit_path, data_only=False)
    require(workbook.sheetnames == SHEETS, f"Lab {lab_number} Audit Log sheet order differs.")
    metadata = workbook[SHEETS[0]]
    require(metadata_value(metadata, "Student Name:") == "Nguyễn Đình Chương", "Student name mismatch.")
    require(metadata_value(metadata, "Student ID:") == "SE204330", "Student ID mismatch.")
    require(metadata_value(metadata, "Course:") == "LAB211", "Course mismatch.")
    require(metadata_value(metadata, "Assignment:") == expected["assignment"], "Assignment mismatch.")
    require(metadata_value(metadata, "Core Prompts Logged:") == expected["prompts"], "Prompt count mismatch.")
    require(workbook.properties.creator == "Nguyễn Đình Chương", "Audit Log author metadata mismatch.")
    require(all(not sheet._images for sheet in workbook.worksheets), "Unexpected embedded image in Audit Log.")

    evidence = workbook[SHEETS[4]]
    evidence_rows = [
        row
        for row in evidence.iter_rows(values_only=True)
        if isinstance(row[0], str) and row[0].startswith("EV-")
    ]
    require(len(evidence_rows) == expected["prompts"], f"Lab {lab_number} Evidence Index count mismatch.")
    require(
        [row[0] for row in evidence_rows]
        == [f"EV-{index:03d}" for index in range(1, int(expected["prompts"]) + 1)],
        f"Lab {lab_number} Evidence IDs are not sequential.",
    )
    harness_hash = sha256(harness_path)[:16]
    output_hash = sha256(output_path)[:16]
    for row in evidence_rows:
        require(result in str(row[5]), f"Lab {lab_number} evidence result mismatch: {row[0]}")
        require(harness_hash in str(row[7]), f"Lab {lab_number} harness hash mismatch: {row[0]}")
        require(output_hash in str(row[7]), f"Lab {lab_number} output hash mismatch: {row[0]}")

    all_values = "\n".join(workbook_values(workbook))
    require("Java 8/Ant PASS" not in all_values, f"Lab {lab_number} contains stale Ant claim.")
    require("Report Section 4 UML" not in all_values, f"Lab {lab_number} contains stale report section.")
    require(str(expected["code"]) in all_values and result in all_values, f"Lab {lab_number} alignment failed.")

    print(
        f"PASS Lab {lab_number}: PDF 7 pages; DOCX readable; "
        f"Audit Log {len(evidence_rows)} evidence rows; hashes aligned."
    )


def main() -> int:
    try:
        java_count = validate_java_imports()
        print(f"PASS imports: {java_count} Java files; no wildcard or definitely unused explicit imports.")
        for lab_number, expected in LABS.items():
            validate_lab(lab_number, expected)
        checksum_count = validate_release_checksums()
        print(f"PASS integrity: {checksum_count} release checksums match.")
    except (AssertionError, OSError, ValueError) as error:
        print(f"VALIDATION FAILED: {error}", file=sys.stderr)
        return 1
    print("ARTIFACT VALIDATION PASSED: reports, Audit Logs, source evidence, and hashes align.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
